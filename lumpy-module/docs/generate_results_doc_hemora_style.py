"""Builds the Hemora-B-section-style Word version of the LSD results.

Run with the lumpy-module venv (needs python-docx):
    venv/Scripts/python.exe docs/generate_results_doc_hemora_style.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIG_DIR = "figures"

doc = Document()

section = doc.sections[0]
section.top_margin = Mm(15)
section.bottom_margin = Mm(15)
section.left_margin = Mm(14)
section.right_margin = Mm(14)

sectPr = section._sectPr
cols = OxmlElement("w:cols")
cols.set(qn("w:num"), "2")
cols.set(qn("w:space"), "425")
sectPr.append(cols)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(9.5)
style.paragraph_format.space_after = Pt(5)
style.paragraph_format.line_spacing = 1.05


def add_heading_b(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10.5)
    return p


def add_para(text, size=9.5, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def set_cell_shading(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(caption, headers, rows):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(3)
    run = cap.add_run(caption)
    run.bold = True
    run.font.size = Pt(8.5)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "1F4E79")
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(255, 255, 255)

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(path, caption, width_in=3.1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(f"{FIG_DIR}/{path}", width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(8)
    cap.paragraph_format.space_after = Pt(6)


add_heading_b("B. Lumpy Skin Disease (LSD) Detection")

add_para(
    "The YOLOv8s model was used to detect candidate skin nodules in cattle "
    "photographs, while ResNet50 was used to classify each detected region as "
    "LSD-positive or healthy. TABLE IV summarizes the performance achieved by both "
    "models."
)

add_table(
    "TABLE IV. MODEL PERFORMANCE SUMMARY",
    ["Model", "Accuracy"],
    [
        ["YOLOv8s", "11.83% (mAP)"],
        ["ResNet50", "93.10%"],
    ],
)

add_para(
    "As shown in Fig. 1, YOLOv8s box loss decreases steadily across training, with "
    "precision and recall trending upward over the 79 training epochs. On the "
    "held-out test set, the detector achieved an overall detection accuracy (mAP@0.5) "
    "of 11.83%."
)

add_figure(
    "yolo_cell14_out0.png",
    "Fig. 1. YOLOv8s training curves across 79 epochs.",
    width_in=3.3,
)

add_para(
    "The ResNet50 model was trained in two phases, and its training performance is "
    "illustrated in Fig. 2, showing the accuracy and AUC curves. Training accuracy "
    "reaches 100% while validation accuracy stabilises around 95%, and the model "
    "achieves 93.1% accuracy with an AUC of 0.9637 on the held-out test set."
)

add_figure(
    "resnet_cell20_out0.png",
    "Fig. 2. ResNet50 accuracy, loss, and AUC across both training phases.",
    width_in=3.3,
)

add_para(
    "The model correctly classified the majority of the test subset, identifying 6 "
    "of 9 healthy cases and 48 of 49 LSD-positive cases. The confusion matrix in "
    "Fig. 3 confirms strong class separation for the LSD-positive class on this "
    "dataset."
)

add_figure(
    "resnet_cell24_out0.png",
    "Fig. 3. Confusion matrix — ResNet50 LSD classification (test set).",
    width_in=2.6,
)

add_para(
    "Overall, the results show that combining nodule detection and image "
    "classification provides a practical basis for automated LSD screening. The "
    "system further integrates an optional farmer-reported symptom checklist, where "
    "image-based results contribute 70% and symptom-based analysis contributes 30% "
    "to the final decision, improving the overall reliability of the system by "
    "incorporating both visual and clinical information."
)

OUT_NAME = "LSD_Results_Hemora_Style_v2.docx"
doc.save(OUT_NAME)
print(f"saved {OUT_NAME}")
