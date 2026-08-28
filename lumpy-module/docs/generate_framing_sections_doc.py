"""Builds the Word version of the CattleSense joint-paper framing sections
(Abstract, Keywords, Introduction, Lit-Review intro, Methodology intro, Conclusion).

Run with the lumpy-module venv (needs python-docx):
    venv/Scripts/python.exe docs/generate_framing_sections_doc.py
"""
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)
section.left_margin = Mm(20)
section.right_margin = Mm(20)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.2


def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    return p


def add_section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return p


def add_label_para(label, text, size=9.5):
    """IEEE-style 'Abstract—...' / 'Keywords—...' bold-italic lead-in paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(f"{label}—")
    r1.bold = True
    r1.italic = True
    r1.font.size = Pt(size)
    r2 = p.add_run(text)
    r2.bold = True
    r2.italic = True
    r2.font.size = Pt(size)
    return p


def add_body_para(text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.first_line_indent = Pt(14)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def add_wordcount(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ---- Title ------------------------------------------------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("CattleSense: Framing Sections for the Joint Paper")
r.bold = True
r.font.size = Pt(15)

add_note(
    "These are the pieces that wrap around your four already-written A–D "
    "subsections (Mastitis, FMD, LSD, Milk Fever) — Abstract, Keywords, "
    "Introduction, a short Literature Review intro, a short Methodology intro, "
    "and a Conclusion. Drop each into its place in the paper; your existing "
    "per-component content stays exactly as your team wrote it."
)

# ---- Abstract -----------------------------------------------------------------
add_section_heading("Abstract")
add_label_para(
    "Abstract",
    "Cattle health disorders — Mastitis, Foot-and-Mouth Disease (FMD), Lumpy "
    "Skin Disease (LSD), and Milk Fever (hypocalcaemia) — pose significant "
    "challenges to Sri Lanka's dairy industry, where diagnostic delays and "
    "limited veterinary access in rural farming communities often result in "
    "late intervention, reduced milk yield, and preventable economic loss. "
    "This research introduces “CattleSense,” an integrated multimodal "
    "machine learning platform that automates early screening for all four "
    "conditions through a single farmer-accessible web application. Each "
    "disease module combines complementary data sources through its own "
    "hybrid fusion strategy: Mastitis detection pairs a ResNet50 udder-image "
    "classifier with a Decision Tree model trained on farmer-collectible milk "
    "biomarkers; FMD detection uses a fine-tuned EfficientNetB0 lesion "
    "classifier alongside a weather-derived outbreak-risk signal; LSD "
    "detection chains a YOLOv8s nodule detector with a ResNet50 classifier "
    "and an optional clinical symptom checklist; and Milk Fever risk is "
    "predicted by an ensemble of Random Forest and XGBoost models trained on "
    "farmer-observable and IoT sensor-derived indicators, adjusted for "
    "heat-stress risk via a live Temperature-Humidity Index. Trained and "
    "evaluated on datasets reflecting Sri Lankan farm conditions, the system "
    "achieved a held-out test accuracy of 87.0% (AUC 0.94) for Mastitis image "
    "classification and 99.4% for its biomarker model, 94.1% (AUC 0.97) for "
    "FMD classification, 93.1% (AUC 0.96) for LSD classification, and 97.0% "
    "(AUC 0.999) for Milk Fever stage prediction. Each module returns a risk "
    "level and actionable, farmer-friendly guidance without requiring "
    "laboratory access or a veterinary login. These results demonstrate that "
    "combining visual, physiological, and environmental evidence through "
    "disease-specific hybrid fusion produces a practical, accessible "
    "early-warning platform capable of reducing diagnostic delays and "
    "supporting timelier veterinary intervention across Sri Lankan dairy "
    "farms.",
)
add_wordcount("≈240 words — trim the mid-sentence metric list if your venue has a hard abstract word cap.")

# ---- Keywords -----------------------------------------------------------------
add_label_para(
    "Keywords",
    "Mastitis, Foot-and-Mouth Disease, Lumpy Skin Disease, Milk Fever, "
    "Multimodal Deep Learning, Hybrid Fusion, Precision Livestock Farming, "
    "Cattle Disease Detection",
)

# ---- Introduction -------------------------------------------------------------
add_section_heading("I. Introduction")

add_body_para(
    "Cattle farming is a cornerstone of Sri Lanka's rural economy, with "
    "approximately 1.2 million cattle supporting dairy production, "
    "agricultural labour, and household income for smallholder farmers "
    "across the country. This livelihood is repeatedly threatened by a small "
    "set of recurring diseases — Mastitis, Foot-and-Mouth Disease, Lumpy Skin "
    "Disease, and Milk Fever — each of which, left undetected until clinical "
    "signs are advanced, results in reduced milk yield, treatment costs, "
    "animal welfare harm, and in the case of the two infectious diseases, "
    "herd-wide and cross-farm transmission."
)

add_body_para(
    "Current diagnostic practice for all four conditions still relies "
    "predominantly on manual observation: a farmer or veterinarian visually "
    "inspecting udders for swelling, skin for nodules, the mouth and hooves "
    "for lesions, or a freshly-calved cow for signs of weakness. This is "
    "inherently subjective, dependent on the examiner's experience, and "
    "frequently delayed — veterinary officers and diagnostic laboratories are "
    "unevenly distributed between urban and rural regions, and smallholder "
    "farmers in outlying areas routinely wait days for a professional "
    "assessment that early-stage disease does not afford. Foot-and-Mouth "
    "Disease and Lumpy Skin Disease compound this problem further: both are "
    "highly contagious, spreading rapidly through direct contact or insect "
    "vectors, so a delay in flagging even one animal can escalate into a "
    "herd- or district-level outbreak."
)

add_body_para(
    "Machine learning offers a complementary path to earlier, more "
    "consistent screening. Convolutional neural networks have demonstrated "
    "strong performance identifying disease-specific visual signatures — "
    "udder inflammation, oral and pedal lesions, skin nodules — while "
    "classical and ensemble machine learning models have proven equally "
    "capable of learning from the structured clinical and behavioural "
    "indicators (milk biomarkers, body condition, calving history, activity "
    "level) that farmers can plausibly observe or measure themselves without "
    "laboratory access. Neither modality alone is complete: image-based "
    "models can miss disease that has not yet produced a visible sign, and "
    "purely numerical models cannot see the lesion a photograph can. This "
    "motivates a hybrid, multimodal approach for each disease individually, "
    "rather than a single one-size-fits-all classifier."
)

add_body_para(
    "This paper presents CattleSense, a four-module early-detection platform "
    "developed collaboratively, with each author responsible for one "
    "disease-specific module — Mastitis, FMD, LSD, and Milk Fever — sharing a "
    "common web-based delivery mechanism, authentication layer, and "
    "cattle-record database. Each module is described independently in "
    "Sections II–IV (Literature Review, Methodology, and Results and "
    "Discussion, respectively), followed by a joint discussion of what the "
    "four modules demonstrate together about hybrid, farmer-accessible "
    "disease screening in a resource-limited dairy farming context."
)
add_wordcount("≈390 words")

# ---- Literature Review intro ---------------------------------------------------
add_section_heading("II. Literature Review — Introductory Paragraph")
add_note("Place this immediately before your four existing A–D literature-review subsections.")

add_body_para(
    "The advancement of automated diagnostic systems for livestock disease "
    "detection has been extensively documented in recent research, with a "
    "growing emphasis on deep learning architectures and multimodal data "
    "fusion combining visual, clinical, and environmental signals. Manual "
    "veterinary inspection has been repeatedly identified as a bottleneck in "
    "dairy herd health management, particularly in resource-limited rural "
    "settings such as Sri Lanka, where veterinary access is scarce and "
    "diagnostic delays are common across all four conditions studied here. "
    "This review categorizes existing research into the four specific "
    "diseases targeted by the CattleSense framework: Mastitis, "
    "Foot-and-Mouth Disease, Lumpy Skin Disease, and Milk Fever."
)
add_wordcount("≈85 words")

# ---- Methodology intro ---------------------------------------------------------
add_section_heading("III. Methodology — Introductory Paragraph")
add_note("Place this immediately before your four existing A–D methodology subsections.")

add_body_para(
    "Each of the four disease modules was developed and evaluated "
    "independently, reflecting the different nature of evidence available "
    "for each condition — image data alone is sufficient for some, while "
    "others depend more heavily on clinical, behavioural, or environmental "
    "indicators. Despite these differences, all four modules follow a shared "
    "design pattern: a primary machine learning or deep learning model "
    "trained on the strongest available signal for that disease, combined "
    "through a disease-appropriate hybrid fusion layer with a secondary, "
    "farmer-collectible data source, and evaluated on a held-out test split "
    "never seen during training. Table numbering and figure numbering below "
    "follow the order Mastitis (A), FMD (B), LSD (C), and Milk Fever (D)."
)
add_wordcount("≈115 words")

# ---- Conclusion -----------------------------------------------------------------
add_section_heading("V. Conclusion")

add_body_para(
    "This study presents CattleSense, a multimodal machine learning platform "
    "for the early detection of four major cattle diseases affecting Sri "
    "Lankan dairy farms — Mastitis, Foot-and-Mouth Disease, Lumpy Skin "
    "Disease, and Milk Fever. By pairing image-based deep learning with "
    "disease-appropriate clinical, biomarker, or environmental data through "
    "hybrid fusion, each module achieved reliable held-out test performance: "
    "87.0% accuracy (AUC 0.94) for Mastitis image classification alongside "
    "99.4% for its biomarker model, 94.1% accuracy (AUC 0.97) for FMD, 93.1% "
    "accuracy (AUC 0.96) for LSD classification, and 97.0% accuracy (AUC "
    "0.999) for Milk Fever stage prediction. These results confirm that "
    "combining visual evidence with the physiological, behavioural, and "
    "environmental data a farmer can realistically supply meaningfully "
    "improves diagnostic reliability over any single modality, while keeping "
    "the system usable without laboratory access or a veterinary login."
)

add_body_para(
    "CattleSense addresses a recurring gap identified across all four "
    "literature reviews: the absence of farmer-accessible, Sri Lanka-specific "
    "automated screening tools for these conditions. By delivering a risk "
    "level and actionable guidance directly to the farmer within seconds of "
    "a photograph or a short symptom checklist, the platform is positioned "
    "to reduce diagnostic delays, support earlier veterinary consultation, "
    "and limit the spread of the two contagious diseases in particular. "
    "Future work will focus on expanding each module's training dataset — "
    "most notably the LSD nodule-detection stage, whose smaller image set "
    "currently limits detection recall — integrating longitudinal per-animal "
    "risk tracking across all four modules, and conducting field trials with "
    "participating farms to validate real-world performance beyond the "
    "held-out test splits reported here."
)
add_wordcount("≈240 words")

doc.save("CattleSense_Paper_Framing_Sections.docx")
print("saved CattleSense_Paper_Framing_Sections.docx")
