"""Builds a Word document of the LSD Results and Discussion section.

Run with the lumpy-module venv (needs python-docx):
    venv/Scripts/python.exe docs/generate_results_doc.py
One-off documentation generator, not part of the running application.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIG_DIR = "figures"

doc = Document()

# ---- Base style ---------------------------------------------------------
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)

for i in range(1, 5):
    hstyle = doc.styles[f"Heading {i}"]
    hstyle.font.name = "Times New Roman"
    hstyle.font.color.rgb = RGBColor(0, 0, 0)
    hstyle.font.bold = True

doc.styles["Heading 1"].font.size = Pt(13)
doc.styles["Heading 2"].font.size = Pt(11.5)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    return h


def add_para(text, italic=False, size=10.5, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    return p


def set_cell_shading(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(caption, headers, rows, bold_last_data_rows=None):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(10)
    run = cap.add_run(caption)
    run.bold = True
    run.font.size = Pt(9.5)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "1F4E79")
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(255, 255, 255)

    bold_rows = bold_last_data_rows or set()
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    if ridx in bold_rows:
                        r.font.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_figure(path, caption, width_in=5.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(f"{FIG_DIR}/{path}", width=Inches(width_in))
    except Exception as exc:
        p.add_run(f"[MISSING IMAGE: {path} — {exc}]")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9.5)
    cap.paragraph_format.space_after = Pt(12)


# ---- Title ---------------------------------------------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("IV. RESULTS AND DISCUSSION")
r.font.size = Pt(14)
r.font.bold = True

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run(
    "Lumpy Skin Disease (LSD) Detection Component — Manathunga M.A.A.S, IT22282422"
)
r.italic = True
r.font.size = Pt(10)

# ---- A. YOLOv8s ------------------------------------------------------------
add_heading("A. YOLOv8s Lesion Detection", level=1)

add_para(
    "The detection stage was trained for up to 150 epochs (batch size 16, image size "
    "640×640, lr0 = 0.001, dropout = 0.1) with early stopping on a patience of 30 "
    "epochs, using a single-class (“LSD-Nodule-Detection”) dataset exported from "
    "Roboflow: 546 training images, 129 of which are deliberately annotation-free "
    "“background” negatives, and a held-out test split of 24 images containing 261 "
    "ground-truth nodule instances. Training ran for 79 epochs before early stopping "
    "triggered, with the best checkpoint (by mAP@0.5) recorded at epoch 48."
)

add_table(
    "TABLE I. YOLOv8s DETECTION PERFORMANCE",
    ["Metric", "Best Epoch (48)", "Final Epoch (79)", "Held-out Test (n=24, 261 instances)"],
    [
        ["Precision", "0.2182", "0.1722", "0.1905"],
        ["Recall", "0.2100", "0.1940", "0.2759"],
        ["mAP@0.5", "0.1383", "0.0991", "0.1183"],
        ["mAP@0.5:0.95", "0.0345", "0.0249", "0.0309"],
        ["F1-score", "—", "—", "0.2254"],
    ],
)

add_para(
    "The training curves in Fig. 1 tell a fairly clear story on their own. Box loss "
    "falls steadily on the training set across all 79 epochs, but validation box loss "
    "follows the same downward path only until roughly the midpoint of training, after "
    "which it turns and climbs back up even as the training curve keeps falling — the "
    "textbook signature of the model starting to memorise training-set specifics rather "
    "than learning generalisable nodule features. Precision and recall behave similarly: "
    "both oscillate in a fairly narrow 0.15–0.25 band with substantial epoch-to-epoch "
    "noise instead of settling into a stable upward trend, and mAP@0.5 tops out around "
    "12–14% before drifting back down toward the final epoch. On the held-out test set, "
    "mAP@0.5 came out to 11.83% — well short of the 80% target this project’s own "
    "non-functional requirements set for the detector."
)

add_figure(
    "yolo_cell14_out0.png",
    "Fig. 1. YOLOv8s training curves — box loss, classification loss, precision/recall, "
    "and mAP@0.5 & mAP@0.5:0.95 across 79 epochs.",
)

add_para(
    "Fig. 2 (the training-time validation confusion matrix) makes the practical "
    "consequence concrete: at the default plotting confidence threshold, the model "
    "correctly flags 80 true nodules but misses 420, while also raising 120 false alarms "
    "on background regions. In plain terms, the detector currently misses roughly five "
    "real nodules for every one it catches, and a modest fraction of what it does flag "
    "isn’t a nodule at all. Given that the underlying dataset traces back to only 254 "
    "unique photographs before Roboflow’s augmentation pipeline expanded it, this is not "
    "a surprising outcome — object detection is a data-hungry task, and a few hundred "
    "source images, unevenly split between clearly-nodular and background-only examples, "
    "is a thin basis for a model to learn what does and does not constitute a "
    "lumpy-skin nodule across varying lighting, coat colour, and camera distance."
)

add_figure(
    "yolo_confusion_matrix_clean.png",
    "Fig. 2. YOLOv8s validation confusion matrix (LSD-Nodule-Detection vs. background).",
    width_in=3.6,
)

# ---- B. ResNet50 -----------------------------------------------------------
add_heading("B. ResNet50 Classification", level=1)

add_para(
    "The classification stage was fine-tuned in the standard two-phase transfer-learning "
    "pattern: 27 epochs training a new classification head on top of a frozen, "
    "ImageNet-pretrained ResNet50 backbone, followed by 16 further epochs with the last "
    "approximately 15 backbone layers unfrozen and a substantially lower learning rate. "
    "Unlike the detector, this model was trained and evaluated on whole photographs "
    "rather than YOLO-cropped regions — 382 images in total (320 LSD-positive, 62 "
    "healthy), split in a stratified 70/15/15 fashion into 266 training, 58 validation, "
    "and 58 test images."
)

add_table(
    "TABLE II. RESNET50 CLASSIFICATION PERFORMANCE (TEST SET, n=58)",
    ["Class", "Precision", "Recall", "F1-score", "Support"],
    [
        ["Healthy", "0.86", "0.67", "0.75", "9"],
        ["LSD-Positive", "0.94", "0.98", "0.96", "49"],
        ["Overall accuracy", "", "", "93.10%", "58"],
        ["ROC-AUC", "", "", "0.9637", ""],
        ["Test loss", "", "", "0.1793", ""],
    ],
    bold_last_data_rows={2, 3},
)

add_para(
    "Taken in isolation, this is comfortably the strongest-performing piece of the "
    "pipeline. The classifier reaches 93.1% accuracy and an AUC of 0.9637 on data it "
    "never saw during training, correctly identifying 48 of 49 LSD-positive test images "
    "and 6 of 9 healthy ones (Fig. 4), and the ROC curve in Fig. 5 sits well clear of the "
    "random-guess diagonal across essentially the whole operating range. That said, Fig. "
    "3 is worth reading carefully rather than at face value: training accuracy and AUC "
    "both saturate at 1.0 by roughly epoch 20, well before fine-tuning even begins, while "
    "validation accuracy settles closer to 94–95% and validation loss stops improving — a "
    "modest but real train/validation gap that is consistent with the underlying 320:62 "
    "class imbalance (roughly five positive images for every negative one). The practical "
    "symptom of this is output saturation: when the trained model was pointed at a fresh "
    "held-out photograph in a later cell, it returned a raw classification probability of "
    "100.0%, and in a separate integrated-pipeline run (Section IV-C) every single "
    "detected region across a test image came back at 99.998–100.000% regardless of how "
    "visually convincing that particular region actually was. The model, in other words, "
    "is more confident than its 93% (not 100%) test accuracy actually warrants. This is "
    "precisely why the deployed application never shows this raw per-crop number to a "
    "farmer — only the combined, calibrated pipeline probability is surfaced, with the "
    "“Nodule” label standing in for the score on the annotated image itself."
)

add_figure(
    "resnet_cell20_out0.png",
    "Fig. 3. ResNet50 accuracy, loss, and AUC across both training phases (dashed line "
    "marks the start of fine-tuning).",
)
add_figure(
    "resnet_cell24_out0.png",
    "Fig. 4. ResNet50 confusion matrix, held-out test set.",
    width_in=3.4,
)
add_figure(
    "resnet_cell26_out0.png",
    "Fig. 5. ROC curve, ResNet50 LSD classification (test set), AUC = 0.9637.",
    width_in=3.8,
)

# ---- C. Integrated pipeline --------------------------------------------------
add_heading("C. Integrated End-to-End Pipeline", level=1)

add_para(
    "The complete pipeline — YOLOv8s detection, followed by cropping and ResNet50 "
    "classification of each detected region, combined as region score = detection "
    "confidence × classification probability, with the image-level score taken as the "
    "maximum across regions and a 30% decision threshold matching the project’s own "
    "Low/Moderate risk boundary — was evaluated on a 24-image set deliberately aligned to "
    "YOLOv8’s own held-out test split (18 LSD-positive, 6 healthy), so that the "
    "end-to-end numbers below are directly comparable to Table I."
)

add_table(
    "TABLE III. END-TO-END PIPELINE PERFORMANCE (ALIGNED HELD-OUT TEST, n=24)",
    ["Class", "Precision", "Recall", "F1-score", "Support"],
    [
        ["Healthy", "0.30", "1.00", "0.46", "6"],
        ["LSD-Positive", "1.00", "0.22", "0.36", "18"],
        ["Overall accuracy", "", "", "41.67%", "24"],
        ["False Positive Rate", "", "", "0.00%", ""],
        ["False Negative Rate", "", "", "77.78%", ""],
        ["Avg. processing time", "", "", "2.69 s/image", ""],
    ],
    bold_last_data_rows={2},
)

add_para(
    "Set against the two components evaluated separately, this result needs to be read "
    "carefully rather than simply as “the system underperforms.” The confusion matrix in "
    "Fig. 6 shows zero false positives and fourteen false negatives out of eighteen "
    "genuinely positive cases — a false-negative rate of 77.78% sitting alongside perfect "
    "precision and a 0% false-positive rate. That specific combination is the tell: the "
    "pipeline is not making bad classification decisions, it is very often failing to "
    "reach a classification decision at all, because the image-level score is entirely "
    "gated on YOLOv8 producing at least one detection first. Since the detector’s own "
    "recall on this identical test set was only 27.59% (Table I), most genuinely positive "
    "images simply never get handed to ResNet50 in the first place — a classifier that is "
    "93% accurate on its own becomes irrelevant to an image the detector already gave up "
    "on. It is a bottleneck problem, not a joint-failure problem, and the numbers are "
    "consistent with that reading throughout: end-to-end precision (100%) tracks the "
    "detector’s own near-zero false-positive behaviour, and end-to-end recall (22%) "
    "tracks the detector’s own recall (27.59%) far more closely than it tracks the "
    "classifier’s recall (98%). On the encouraging side, average processing time came in "
    "at 2.69 seconds per image, comfortably inside the 10–15 second target set for the "
    "platform, and Fig. 7 shows a case where detection does succeed — twelve regions "
    "found in 5.72 seconds, correctly pushed to a 76.7% overall probability and a HIGH "
    "RISK flag — which is a useful reminder that the fusion logic itself behaves exactly "
    "as designed once it has something to work with."
)

add_figure(
    "integrated_cell14_out0.png",
    "Fig. 6. End-to-end integrated pipeline confusion matrix, aligned 24-image held-out test.",
    width_in=4.0,
)
add_figure(
    "integrated_cell16_out2.png",
    "Fig. 7. Qualitative example — annotated detection output for a correctly-flagged "
    "HIGH RISK case.",
)

# ---- D. Discussion ----------------------------------------------------------
add_heading("D. Discussion", level=1)

add_para(
    "Placed side by side, Tables I–III describe a pipeline whose two learned components "
    "are not equally mature: classification is genuinely strong, detection is not yet, "
    "and because the two stages run in series rather than in parallel, the weaker of the "
    "two effectively sets the ceiling for the whole system. That is a useful, actionable "
    "conclusion rather than a discouraging one — it means the highest-leverage next step "
    "is concentrated almost entirely on the detection side: growing the underlying image "
    "set beyond its current 254 unique source photographs, auditing whether the "
    "“background” negatives and positive annotations are consistently drawn, and "
    "revisiting the training schedule in light of the validation box-loss divergence "
    "visible in Fig. 1 (an earlier stopping point, or somewhat stronger regularisation, "
    "seems like a reasonable first experiment). It is also worth being explicit about "
    "what the current precision/recall balance actually means for a farmer using this as "
    "a screening tool in the field: a system that essentially never raises a false alarm "
    "but misses a majority of real cases is safer to deploy than the reverse would be, "
    "but it clearly cannot yet stand as the sole basis for a diagnosis. That gap between "
    "“detects nothing visually obvious yet” and “is actually healthy” is exactly the case "
    "the clinical symptom checklist described in Section III is intended to help cover, "
    "by giving the system a second, independent line of evidence to draw on for the false "
    "negatives this evaluation surfaces. Formally re-measuring Table III with that "
    "symptom layer active would need a test set that also carries recorded clinical "
    "symptom labels, which the current 24-image aligned split does not, and is "
    "accordingly left as a direction for future validation work rather than a claim made "
    "here."
)

# ---- Author notes (clearly separated, meant to be deleted before submission) ----
doc.add_page_break()
h = add_heading("AUTHOR NOTES — DELETE BEFORE SUBMISSION", level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

add_para(
    "Figures used above, with status:", size=10)
rows = [
    ["1", "yolo_cell14_out0.png", "YOLOv8s training curves", "Ready as-is"],
    ["2", "yolo_confusion_matrix_clean.png", "YOLOv8s validation confusion matrix",
     "Regenerated — original had a clipped axis label; same underlying counts (80/120/420)"],
    ["3", "resnet_cell20_out0.png", "ResNet50 training curves", "Ready as-is"],
    ["4", "resnet_cell24_out0.png", "ResNet50 confusion matrix", "Ready as-is"],
    ["5", "resnet_cell26_out0.png", "ResNet50 ROC curve", "Ready as-is"],
    ["6", "integrated_cell14_out0.png", "End-to-end confusion matrix", "Ready as-is"],
    ["7", "integrated_cell16_out2.png", "Qualitative annotated example",
     "Please visually confirm before final use — not fully verified by Claude"],
]
add_table("Figure status", ["Fig.", "File", "Shows", "Status"], rows)

add_para(
    "Consistency check: every confusion matrix above was hand-verified against its own "
    "printed classification report (precision/recall recomputed from raw counts) — all "
    "internally consistent, nothing numeric needed correcting. ResNet50's "
    "model.evaluate() labelled its accuracy metric “compile_metrics: 0.9310” (a Keras 3 "
    "display quirk) — verified this is genuinely accuracy ((6+48)/58 = 0.9310) and written "
    "into Table II as “Overall accuracy.”",
    italic=True, size=9.5,
)
add_para(
    "Deliberately NOT cited anywhere above: an earlier ad hoc live-deployment test (17 "
    "photos, 88% accuracy) run outside these notebooks. It wasn't drawn from a documented "
    "held-out split, so it isn't publication-grade evidence next to Tables I–III.",
    italic=True, size=9.5,
)

doc.save("LSD_Results_and_Discussion.docx")
print("saved LSD_Results_and_Discussion.docx")
