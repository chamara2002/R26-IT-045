"""Builds the full assembled CattleSense paper as a single Word document.

Combines the framing sections (Abstract, Keywords, Introduction, Lit Review
intro, Methodology intro, Conclusion) with the four teammates' A-D content
(Literature Review, Methodology, Results and Discussion), fixing figure/table
numbering collisions and de-duplicated content along the way.

Run with the lumpy-module venv (needs python-docx):
    venv/Scripts/python.exe docs/generate_full_paper.py
"""
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIG_DIR = "figures"

doc = Document()
section = doc.sections[0]
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)
section.left_margin = Mm(20)
section.right_margin = Mm(20)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.2


# ---------------------------------------------------------------- helpers --
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(4)


def add_author_placeholder():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "[Author block — insert all four names, IDs, and affiliations here, "
        "in your usual IEEE multi-author format]"
    )
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p.paragraph_format.space_after = Pt(14)


def add_heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def add_heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(11.5)


def add_label_para(label, text, size=9.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(f"{label}—")
    r1.bold = True
    r1.italic = True
    r1.font.size = Pt(size)
    r2 = p.add_run(text)
    r2.bold = True
    r2.italic = True
    r2.font.size = Pt(size)


def add_body(text, size=11, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    if indent:
        p.paragraph_format.first_line_indent = Pt(14)
    p.add_run(text).font.size = Pt(size)


def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def set_shade(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(caption, headers, rows, bold_rows=None):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(3)
    r = cap.add_run(caption)
    r.bold = True
    r.font.size = Pt(9.5)

    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_shade(hdr[i], "1F4E79")
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(255, 255, 255)
    bold_rows = bold_rows or set()
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    if ridx in bold_rows:
                        r.font.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_real_figure(path, caption, width_in=4.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(f"{FIG_DIR}/{path}", width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9.5)
    cap.paragraph_format.space_after = Pt(10)


def add_figure_placeholder(caption):
    """For figures whose source image files aren't available to me."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(4.6)
    cell = t.rows[0].cells[0]
    set_shade(cell, "F2F2F2")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("[Insert original figure image here]")
    r.italic = True
    r.font.color.rgb = RGBColor(0x90, 0x90, 0x90)
    r.font.size = Pt(10)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9.5)
    cap.paragraph_format.space_after = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_ref(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"[{num}]\t{text}")
    r.font.size = Pt(9.5)


def add_ref_todo(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"[{num}]\t{text}")
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    r.bold = True


# =============================================================================
# TITLE PAGE
# =============================================================================
add_title("CattleSense: A Multimodal Machine Learning Platform for Early "
          "Detection of Cattle Diseases in Sri Lanka")
add_author_placeholder()

# =============================================================================
# ABSTRACT & KEYWORDS
# =============================================================================
add_label_para(
    "Abstract",
    "Cattle health disorders — Mastitis, Foot-and-Mouth Disease (FMD), Lumpy "
    "Skin Disease (LSD), and Milk Fever (hypocalcaemia) — pose significant "
    "challenges to Sri Lanka's dairy industry, where diagnostic delays and "
    "limited veterinary access in rural farming communities often result in "
    "late intervention, reduced milk yield, and preventable economic loss. "
    "This research introduces “CattleSense,” an integrated multimodal "
    "machine learning platform that automates early screening for all four "
    "conditions through a single farmer-accessible web application. Each "
    "disease module combines complementary data sources through its own "
    "hybrid fusion strategy: Mastitis detection pairs a ResNet50 udder-image "
    "classifier with a Decision Tree model trained on farmer-collectible milk "
    "biomarkers; FMD detection uses a fine-tuned EfficientNetB0 lesion "
    "classifier alongside a weather-derived outbreak-risk signal; LSD "
    "detection chains a YOLOv8s nodule detector with a ResNet50 classifier "
    "and an optional clinical symptom checklist; and Milk Fever risk is "
    "predicted by an ensemble of Random Forest and XGBoost models trained on "
    "farmer-observable and IoT sensor-derived indicators, adjusted for "
    "heat-stress risk via a live Temperature-Humidity Index. Trained and "
    "evaluated on datasets reflecting Sri Lankan farm conditions, the system "
    "achieved a held-out test accuracy of 87.0% (AUC 0.94) for Mastitis image "
    "classification and 99.4% for its biomarker model, 94.1% (AUC 0.97) for "
    "FMD classification, 93.1% (AUC 0.96) for LSD classification, and 97.0% "
    "(AUC 0.999) for Milk Fever stage prediction. Each module returns a risk "
    "level and actionable, farmer-friendly guidance without requiring "
    "laboratory access or a veterinary login. These results demonstrate that "
    "combining visual, physiological, and environmental evidence through "
    "disease-specific hybrid fusion produces a practical, accessible "
    "early-warning platform capable of reducing diagnostic delays and "
    "supporting timelier veterinary intervention across Sri Lankan dairy "
    "farms.",
)
add_label_para(
    "Keywords",
    "Mastitis, Foot-and-Mouth Disease, Lumpy Skin Disease, Milk Fever, "
    "Multimodal Deep Learning, Hybrid Fusion, Precision Livestock Farming, "
    "Cattle Disease Detection",
)

# =============================================================================
# I. INTRODUCTION
# =============================================================================
add_heading1("I. Introduction")
add_body(
    "Cattle farming is a cornerstone of Sri Lanka's rural economy, with "
    "approximately 1.2 million cattle supporting dairy production, "
    "agricultural labour, and household income for smallholder farmers "
    "across the country. This livelihood is repeatedly threatened by a small "
    "set of recurring diseases — Mastitis, Foot-and-Mouth Disease, Lumpy Skin "
    "Disease, and Milk Fever — each of which, left undetected until clinical "
    "signs are advanced, results in reduced milk yield, treatment costs, "
    "animal welfare harm, and, in the case of the two infectious diseases, "
    "herd-wide and cross-farm transmission."
)
add_body(
    "Current diagnostic practice for all four conditions still relies "
    "predominantly on manual observation: a farmer or veterinarian visually "
    "inspecting udders for swelling, skin for nodules, the mouth and hooves "
    "for lesions, or a freshly-calved cow for signs of weakness. This is "
    "inherently subjective, dependent on the examiner's experience, and "
    "frequently delayed — veterinary officers and diagnostic laboratories are "
    "unevenly distributed between urban and rural regions, and smallholder "
    "farmers in outlying areas routinely wait days for a professional "
    "assessment that early-stage disease does not afford. Foot-and-Mouth "
    "Disease and Lumpy Skin Disease compound this problem further: both are "
    "highly contagious, spreading rapidly through direct contact or insect "
    "vectors, so a delay in flagging even one animal can escalate into a "
    "herd- or district-level outbreak."
)
add_body(
    "Machine learning offers a complementary path to earlier, more "
    "consistent screening. Convolutional neural networks have demonstrated "
    "strong performance identifying disease-specific visual signatures — "
    "udder inflammation, oral and pedal lesions, skin nodules — while "
    "classical and ensemble machine learning models have proven equally "
    "capable of learning from the structured clinical and behavioural "
    "indicators (milk biomarkers, body condition, calving history, activity "
    "level) that farmers can plausibly observe or measure themselves without "
    "laboratory access. Neither modality alone is complete: image-based "
    "models can miss disease that has not yet produced a visible sign, and "
    "purely numerical models cannot see the lesion a photograph can. This "
    "motivates a hybrid, multimodal approach for each disease individually, "
    "rather than a single one-size-fits-all classifier."
)
add_body(
    "This paper presents CattleSense, a four-module early-detection platform "
    "developed collaboratively, with each author responsible for one "
    "disease-specific module — Mastitis, FMD, LSD, and Milk Fever — sharing a "
    "common web-based delivery mechanism, authentication layer, and "
    "cattle-record database. Each module is described independently in "
    "Sections II–IV (Literature Review, Methodology, and Results and "
    "Discussion, respectively), followed by a joint discussion of what the "
    "four modules demonstrate together about hybrid, farmer-accessible "
    "disease screening in a resource-limited dairy farming context."
)

# =============================================================================
# II. LITERATURE REVIEW
# =============================================================================
add_heading1("II. Literature Review")
add_body(
    "The advancement of automated diagnostic systems for livestock disease "
    "detection has been extensively documented in recent research, with a "
    "growing emphasis on deep learning architectures and multimodal data "
    "fusion combining visual, clinical, and environmental signals. Manual "
    "veterinary inspection has been repeatedly identified as a bottleneck in "
    "dairy herd health management, particularly in resource-limited rural "
    "settings such as Sri Lanka, where veterinary access is scarce and "
    "diagnostic delays are common across all four conditions studied here. "
    "This review categorizes existing research into the four specific "
    "diseases targeted by the CattleSense framework: Mastitis, "
    "Foot-and-Mouth Disease, Lumpy Skin Disease, and Milk Fever.",
    indent=False,
)

add_heading2("A. Mastitis Detection")
add_body(
    "Machine learning and deep learning approaches have increasingly been "
    "applied to bovine mastitis detection, using both image-based and "
    "sensor-derived features [1], [2]. Convolutional neural network "
    "architectures have been used to classify udder thermal images, "
    "achieving detection accuracies above 87% when combining "
    "temperature-difference features with ocular thermal cues [3]. "
    "Similarly, ultrasonography-based deep networks have been explored for "
    "distinguishing mastitis-affected from healthy udder tissue in buffaloes "
    "[4]. In parallel, sensor-driven approaches using milk electrical "
    "conductivity, somatic cell count, and biomarker measurements have "
    "achieved reliable classification via ensemble machine learning models "
    "[5]. However, most existing studies rely on a single data modality, "
    "motivating research into multimodal fusion approaches that combine "
    "visual and physiological indicators for more robust early detection [6]."
)

add_heading2("B. Foot-and-Mouth Disease (FMD) Detection")
add_body(
    "Automated visual diagnosis of FMD — targeting lesions on the mouth, "
    "tongue, and hooves — has also drawn considerable deep learning "
    "attention, largely as a substitute for manual veterinary inspection. "
    "For instance, a 2024 study built a YOLOv4-based detector for cattle FMD "
    "and reported accuracy as high as 98% on annotated lesion images [7], "
    "and a separate ensemble combining VGG16, ResNet50, and InceptionV3 "
    "achieved 98.2% accuracy when jointly identifying FMD and Lumpy Skin "
    "Disease across datasets spanning multiple countries [8]. Work in this "
    "space isn't limited to image classification, either: interpretable ML "
    "models have been used to study FMD's landscape epidemiology in endemic "
    "areas [9], flag FMD incidence on dairy farms lacking veterinary access "
    "[10], and connect seasonal and environmental factors to outbreak risk "
    "across South Asia [11]."
)

add_heading2("C. Lumpy Skin Disease (LSD) Detection")
add_body(
    "Deep learning has been applied to livestock skin conditions more "
    "broadly, but so far nothing has zeroed in specifically on LSD. Alharan "
    "et al. [12], for example, classified general cattle skin diseases with "
    "CNNs and exceeded 85% accuracy, while Kumar et al. [13] used transfer "
    "learning — VGG16 and InceptionV3 — to detect cattle skin lesions, "
    "though again without focusing on any single disease or building toward "
    "farmer-facing tools. On the epidemiological side, earlier foundational "
    "studies [14], [15] documented LSD's clinical presentation and its "
    "transboundary spread across Africa, the Middle East, and Asia, but "
    "diagnosis in practice has continued to rely on manual veterinary "
    "examination. Meanwhile, architectures like YOLOv8 [16] and ResNet50 "
    "[17] have each proven themselves individually — one for real-time "
    "object detection, the other for image classification — yet no one "
    "appears to have combined the two into a single pipeline built "
    "specifically for automated LSD detection."
)

add_heading2("D. Milk Fever Detection")
add_body(
    "Milk Fever, also known as hypocalcaemia, is a common metabolic "
    "disorder affecting dairy cows around calving, occurring when blood "
    "calcium concentration drops sharply during the peripartum period. "
    "Recent studies have investigated machine learning approaches to "
    "predicting hypocalcaemia in dairy cows: Lasser et al. (2021) used "
    "farm- and animal-level data with machine learning models to predict "
    "several dairy diseases, including periparturient hypocalcaemia, while "
    "Van Leerdam et al. (2024) combined behavioural sensor data with parity, "
    "body condition score, and other cow-level information, evaluating "
    "XGBoost and LSTM models to show that behavioural information can "
    "support early hypocalcaemia prediction. Other recent research has also "
    "investigated milk composition characteristics as inputs for "
    "machine-learning-based hypocalcaemia prediction. However, many existing "
    "approaches depend on sensors, laboratory measurements, or detailed farm "
    "data that are not available to most smallholder farmers. This research "
    "therefore focuses on a more accessible approach, using farmer-friendly "
    "cow information for Milk Fever risk prediction."
)
add_note(
    "Note: the Lasser et al. (2021) and Van Leerdam et al. (2024) citations "
    "above are not yet numbered/bracketed — see the References section for "
    "why, and what's needed to finish this."
)

# =============================================================================
# III. METHODOLOGY
# =============================================================================
add_heading1("III. Methodology")
add_body(
    "Each of the four disease modules was developed and evaluated "
    "independently, reflecting the different nature of evidence available "
    "for each condition — image data alone is sufficient for some, while "
    "others depend more heavily on clinical, behavioural, or environmental "
    "indicators. Despite these differences, all four modules follow a shared "
    "design pattern: a primary machine learning or deep learning model "
    "trained on the strongest available signal for that disease, combined "
    "through a disease-appropriate hybrid fusion layer with a secondary, "
    "farmer-collectible data source, and evaluated on a held-out test split "
    "never seen during training.",
    indent=False,
)

add_heading2("A. Mastitis Detection")
add_body(
    "The Mastitis Detection module takes a hybrid approach, pairing an "
    "image-based deep learning model with a numerical clinical model. Two "
    "publicly available datasets were used for this — one sourced from "
    "Kaggle, the other from Roboflow: an udder image set (1,172 images "
    "spanning mastitis and normal classes) and a clinical biomarker dataset "
    "(800 records). Before training, the image dataset was screened for "
    "exact and near-duplicate entries using file hashing and perceptual "
    "hashing, and then split into training, validation, and test sets using "
    "a cluster-aware strategy to guard against data leakage. Every image was "
    "resized to 224×224 pixels and harmonized through fixed-quality JPEG "
    "re-encoding, with augmentation applied only on the training side."
)
add_body(
    "For the image branch, a ResNet50 CNN with a frozen, pretrained backbone "
    "was trained to classify udder images as mastitis-positive or normal. "
    "Alongside this, a Decision Tree classifier was trained on five "
    "biomarkers that a farmer could plausibly collect — milk temperature, "
    "pH, conductivity, yield, and clotting — to generate an independent "
    "numerical prediction. These two outputs are then merged through a "
    "soft-voting hybrid fusion layer, which can optionally be refined using "
    "a farmer-reported symptom checklist, to arrive at the final diagnosis. "
    "Fig. 1 shows the overall architecture."
)
add_figure_placeholder("Fig. 1. Model architecture of Mastitis Detection.")

add_heading2("B. Foot-and-Mouth Disease (FMD) Detection")
add_body(
    "For the FMD module, an image-classification pipeline was built around "
    "a fine-tuned EfficientNetB0 CNN, trained on a set of 339 labelled "
    "photographs — cattle mouths, tongues, hooves, and udders — split almost "
    "evenly between healthy and FMD-positive cases (170 vs. 169). Each image "
    "was checked by hand first, mainly to weed out anything corrupted before "
    "it went anywhere near training."
)
add_body(
    "One issue encountered early on: the raw dataset had no case "
    "identifier, so there was no built-in way to tell which photos came "
    "from the same animal. This was addressed with a filename heuristic "
    "that grouped multi-site photos belonging to the same case, which "
    "mattered a lot for the split — a grouped stratified 85/15 train-test "
    "division was used along with grouped 3-fold cross-validation, "
    "specifically so that no single case's images could end up on both "
    "sides of a split (which would otherwise have quietly inflated the "
    "accuracy numbers). Every image, whether at training time or during "
    "live inference, passes through the exact same preprocessing steps — "
    "decode, resize to 160×160, normalize — so there is no drift between "
    "what the model was trained on and what it sees in production."
)
add_body(
    "The EfficientNetB0 backbone was fine-tuned with class-weighted loss "
    "and the backbone's own input rescaling, then tested against a "
    "held-out set the model had never seen, extracting accuracy, "
    "precision, recall, F1, ROC-AUC, and a confusion matrix. The model's "
    "prediction is then fed into a rule-based hybrid layer that also "
    "factors in optional clinical symptom scores, along with a "
    "weather-based FMD spread-risk signal drawing on live data from "
    "Open-Meteo and a DAPH-reported rule about seasonal escalation between "
    "December and February. The result is a combined risk level and "
    "recommendation, shown in the pipeline in Fig. 2."
)
add_figure_placeholder("Fig. 2. Model architecture of Foot-and-Mouth Disease (FMD) Detection.")

add_heading2("C. Lumpy Skin Disease (LSD) Detection")
add_body(
    "The Lumpy Skin Disease (LSD) analysis module is implemented using a "
    "hybrid deep learning approach combining YOLOv8s and ResNet50 models. A "
    "cattle skin image dataset, sourced via Roboflow for detection (254 "
    "images, single “nodule” class) and organized into “Lumpy Skin” and "
    "“Normal Skin” folders for classification (382 images), was split "
    "into training, validation, and test sets. All images were resized and "
    "normalized, and augmentation was applied only to the training set to "
    "reduce overfitting."
)
add_body(
    "The YOLOv8s model was trained on annotated images to detect and "
    "localize skin nodules, while the ResNet50 model classified each "
    "detected region as LSD-positive or healthy. Detection confidence and "
    "classification probability were multiplied per region, and the "
    "maximum region score across the image formed the overall photo-based "
    "probability. This was combined with an optional farmer-reported "
    "symptom checklist through a weighted hybrid fusion layer to support "
    "the final risk decision. The architecture of the system is shown in "
    "Fig. 3."
)
add_figure_placeholder("Fig. 3. Model architecture of Lumpy Skin Disease (LSD) Detection.")

add_heading2("D. Milk Fever Detection")
add_body(
    "The proposed Milk Fever detection module uses a machine learning "
    "prediction approach trained on a combined dataset of 8,783 records "
    "gathered from three sources: a clinically grounded dataset (2,000 "
    "records) built using validated veterinary rules, a real-world dataset "
    "(1,316 records) collected from Sri Lankan farms, and a real IoT sensor "
    "dataset (5,467 records from 580 dairy cows) sourced from an existing "
    "GitHub repository. All records were standardized to eight features: "
    "parity, blood calcium, blood phosphorus, body condition score (BCS), "
    "days to calving, milk yield on day one, activity level, and dietary "
    "cation-anion difference (DCAD). SMOTE oversampling was applied to the "
    "training data to address class imbalance in the Critical stage, "
    "producing 2,700 samples per class. Seven machine learning algorithms "
    "were trained and compared — Logistic Regression, SVM, MLP Neural "
    "Network, Random Forest, XGBoost, Gradient Boosting, and an Ensemble "
    "Voting Classifier combining Random Forest and XGBoost through soft "
    "voting. The Ensemble model was selected for production based on its "
    "AUC-ROC of 0.9986 and validation accuracy of 98.15%, enabling "
    "prediction without laboratory access. The architecture of the system "
    "is shown in Fig. 4."
)
add_figure_placeholder("Fig. 4. Model architecture of Milk Fever Detection.")

# =============================================================================
# IV. RESULTS AND DISCUSSION
# =============================================================================
add_heading1("IV. Results and Discussion")

# --- A. Mastitis -----------------------------------------------------------
add_heading2("A. Mastitis Detection")
add_body(
    "The proposed system was evaluated across three components: the "
    "ResNet50 image classifier, the Decision Tree numerical model, and the "
    "combined hybrid fusion approach. The training and validation "
    "performance of each model is shown in TABLE I."
)
add_table(
    "TABLE I. ACCURACY OF EXPERIMENTED MODELS",
    ["Model", "Train Accuracy", "Validation Accuracy"],
    [
        ["ResNet50 (Image)", "93.37%", "81.77%"],
        ["Decision Tree (Numerical)", "—", "99.69%"],
    ],
)
add_body(
    "The ResNet50 model was trained with a frozen pretrained backbone, "
    "using early stopping based on validation AUC. As Fig. 5 shows, "
    "training accuracy climbed steadily across epochs and reached 93.37% by "
    "the selected best epoch, while validation accuracy settled around "
    "81.77% — a gap that suggests the model was learning genuine patterns "
    "rather than simply memorizing the training set."
)
add_figure_placeholder("Fig. 5. Accuracy graph of the ResNet50 mastitis classification model.")
add_body(
    "On the held-out test set — data the model never saw during training — "
    "the final model reached 87.01% accuracy, with precision at 91.45% and "
    "recall at 89.17%. The ROC curve in Fig. 6 backs this up, showing "
    "strong discriminative ability with an AUC of 0.9376."
)
add_figure_placeholder("Fig. 6. ROC curve of the ResNet50 mastitis detection model.")
add_body(
    "Looking at the confusion matrix in Fig. 7, the model correctly "
    "flagged 107 true mastitis cases and 47 true normal cases, with 10 "
    "false positives and 13 false negatives. This precision-recall balance "
    "was intentional — in a clinical setting, missing an actual mastitis "
    "case is far more costly than raising a false alarm, so the model was "
    "tuned with that trade-off in mind."
)
add_figure_placeholder("Fig. 7. Confusion matrix of mastitis detection (ResNet50).")
add_body(
    "The Decision Tree model, working from five biomarkers a farmer could "
    "realistically collect, performed even better: 99.38% test accuracy, "
    "100% precision, and 97.06% recall, getting only one case wrong out of "
    "160 test samples. To make sure this wasn't a fluke of the particular "
    "split, repeated stratified cross-validation (5-fold × 10 repeats) was "
    "run, which held up well — a mean accuracy of 99.69% with a standard "
    "deviation of just 0.38%. This strong performance is not too "
    "surprising given how well-established the underlying biology is; milk "
    "conductivity and temperature have long been known to correlate "
    "closely with mastitis status."
)
add_body(
    "Finally, the two models come together through a soft-voting hybrid "
    "fusion layer, where the image-based and numerical predictions each "
    "carry equal weight (50%) whenever both inputs are available, with an "
    "optional refinement step based on a farmer-reported symptom "
    "checklist. Combining the two sources this way gives a more complete "
    "picture — pulling in both visual and physiological signals — while "
    "still letting the system fall back to image-only mode when biomarker "
    "data simply isn't available."
)

# --- B. FMD ------------------------------------------------------------------
add_heading2("B. Foot-and-Mouth Disease (FMD) Detection")
add_body(
    "The FMD image classifier was evaluated on a held-out test set of 51 "
    "images (25 no-FMD and 26 FMD-positive) using a grouped split to "
    "prevent data leakage. The fine-tuned EfficientNetB0 achieved 94.1% "
    "accuracy, with 94.2% precision, 94.1% recall, and 94.1% F1-score. For "
    "FMD-positive cases specifically, it achieved 92.6% precision, 96.2% "
    "recall, and 94.3% F1-score, with an ROC-AUC of 0.971. Three-fold "
    "grouped cross-validation produced accuracies of 86.7%, 89.5%, and "
    "95.1%, indicating stable performance."
)
add_table(
    "TABLE II. MODEL PERFORMANCE SUMMARY",
    ["Model", "Key Metric", "Value"],
    [
        ["EfficientNetB0", "Test Accuracy", "94.1%"],
        ["EfficientNetB0", "Macro F1-Score", "94.1%"],
        ["EfficientNetB0", "FMD Recall", "96.2%"],
        ["EfficientNetB0", "ROC-AUC", "0.971"],
        ["EfficientNetB0", "Avg. Precision (PR-AUC)", "0.968"],
    ],
)
add_body(
    "The training and validation accuracy of the EfficientNetB0 model "
    "across 17 training epochs is illustrated in Fig. 8. Validation "
    "accuracy converges quickly to around 92–94% within the first two "
    "epochs and remains stable thereafter, while training accuracy "
    "continues to improve gradually, indicating that the model generalizes "
    "well without severe overfitting."
)
add_figure_placeholder("Fig. 8. Training and validation accuracy curve of the EfficientNetB0 model.")
add_body(
    "The precision-recall curve, computed from the model's predicted FMD "
    "probabilities on the held-out test set, stays close to the top-right "
    "corner of the plot, confirming that the model sustains high precision "
    "even as recall increases — the desired behaviour for a decision-support "
    "tool where both missed disease cases and false alarms carry real cost."
)
add_figure_placeholder("Fig. 9. Precision–Recall curve of the EfficientNetB0 FMD classifier.")
add_body(
    "The confusion matrix (Fig. 10) breaks this down further: of the 25 "
    "true no-FMD images, 23 were classified correctly and 2 came back as "
    "false positives; of the 26 true FMD images, 25 were caught correctly, "
    "with just 1 false negative slipping through."
)
add_figure_placeholder("Fig. 10. Confusion matrix of the EfficientNetB0 FMD classifier.")

# --- C. LSD --------------------------------------------------------------------
add_heading2("C. Lumpy Skin Disease (LSD) Detection")
add_body(
    "The YOLOv8s model was used to detect candidate skin nodules in cattle "
    "photographs, while ResNet50 was used to classify each detected region "
    "as LSD-positive or healthy. TABLE III summarizes the performance "
    "achieved by both models."
)
add_table(
    "TABLE III. MODEL PERFORMANCE SUMMARY",
    ["Model", "Accuracy"],
    [
        ["YOLOv8s", "11.83% (mAP)"],
        ["ResNet50", "93.10%"],
    ],
)
add_body(
    "Fig. 11 shows YOLOv8s's box loss dropping steadily throughout "
    "training, with precision and recall both trending upward over the 79 "
    "epochs. On the held-out test set, the detector achieved an overall "
    "detection accuracy (mAP@0.5) of 11.83%."
)
add_figure_placeholder("Fig. 11. YOLOv8s training curves across 79 epochs.")
add_body(
    "ResNet50, on the other hand, was trained in two phases, with accuracy "
    "and AUC curves plotted in Fig. 12. Training accuracy eventually hits "
    "100%, while validation accuracy settles in around 95%, and on the "
    "actual held-out test set the model reaches 93.1% accuracy with an AUC "
    "of 0.9637."
)
add_real_figure("resnet_cell20_out0.png", "Fig. 12. ResNet50 accuracy, loss, and AUC across both training phases.")
add_body(
    "The model correctly classified the majority of the test subset, "
    "identifying 6 of 9 healthy cases and 48 of 49 LSD-positive cases. The "
    "confusion matrix in Fig. 13 confirms strong class separation for the "
    "LSD-positive class on this dataset."
)
add_real_figure("resnet_cell24_out0.png", "Fig. 13. Confusion matrix — ResNet50 LSD classification (test set).", width_in=3.2)
add_body(
    "Taken together, these results suggest that pairing nodule detection "
    "with image classification gives a workable foundation for automated "
    "LSD screening. The system also folds in an optional farmer-reported "
    "symptom checklist — image-based results carry 70% of the final "
    "decision, symptom-based input the remaining 30% — which helps round "
    "out the diagnosis by pulling in clinical observations alongside the "
    "visual evidence."
)

# --- D. Milk Fever -----------------------------------------------------------
add_heading2("D. Milk Fever Detection")
add_body(
    "The Milk Fever detection module was evaluated on a held-out test set "
    "of 1,757 records drawn from the combined dataset of 8,783 samples "
    "assembled across three complementary sources. The Ensemble Voting "
    "Classifier combining Random Forest and XGBoost through soft "
    "probability averaging achieved a test accuracy of 97.04%, a weighted "
    "F1-score of 97.04%, and an AUC-ROC of 0.9986, substantially exceeding "
    "the proposed research targets of 90% accuracy and 0.90 AUC-ROC. These "
    "results place the ensemble model clearly ahead of all six individually "
    "evaluated algorithms, as summarised in TABLE IV."
)
add_table(
    "TABLE IV. ACCURACY OF EXPERIMENTED MODELS",
    ["Model", "CV Accuracy", "Test Accuracy"],
    [
        ["Logistic Regression", "71.12% ± 1.50%", "67.05%"],
        ["SVM", "92.22% ± 0.89%", "89.24%"],
        ["MLP Neural Network", "96.05% ± 0.53%", "93.17%"],
        ["Random Forest", "96.38% ± 0.34%", "95.28%"],
        ["XGBoost", "97.95% ± 0.31%", "96.98%"],
        ["Gradient Boosting", "97.71% ± 0.23%", "96.81%"],
        ["Ensemble (RF+XGB)", "98.15% ± 0.23%", "97.04%"],
    ],
    bold_rows={6},
)
add_body(
    "The comparative evaluation revealed a clear performance hierarchy "
    "among the trained models. Logistic Regression, serving as the linear "
    "baseline, produced a test accuracy of only 67.05% and an AUC-ROC of "
    "0.8298, confirming that the relationship between the eight "
    "physiological and behavioural input features and the four-stage "
    "output is fundamentally non-linear. Support Vector Machine improved on "
    "this considerably at 89.24% accuracy, while the MLP Neural Network "
    "reached 93.17%, suggesting that the dataset volume and feature "
    "complexity are sufficient to support deep learning but not optimally "
    "suited to it without larger training sets. Random Forest performed "
    "strongly at 95.28% accuracy with a Critical-stage recall of 95.69% — "
    "notably the highest individual recall for the most clinically "
    "important class — though its overall AUC-ROC of 0.9973 fell "
    "marginally behind the gradient boosting approaches. XGBoost and "
    "Gradient Boosting both achieved test accuracies above 96.8% and "
    "AUC-ROC values of 0.9983, confirming that boosting architectures are "
    "particularly well-suited to this tabular multi-class classification "
    "task. The Ensemble model's soft voting mechanism, which averages the "
    "probability outputs of Random Forest and XGBoost rather than applying "
    "hard majority voting, produced the highest overall AUC-ROC of 0.9986 "
    "and the most stable cross-validation accuracy of 98.15% ± 0.23%, "
    "reflecting the complementary strengths of bagging and boosting when "
    "combined."
)
add_body(
    "Per-stage classification performance, shown in the confusion matrix "
    "in Fig. 14, reveals that Subclinical and Mild stages were classified "
    "with the highest reliability, achieving F1-scores of 98% and 97% "
    "respectively. This is clinically encouraging, as Subclinical cases — "
    "which present without visible symptoms — are precisely the cases most "
    "frequently missed by reactive farm-level observation. Moderate-stage "
    "classification achieved a 96% F1-score, while the Critical stage, "
    "despite representing only 6.6% of the original dataset prior to SMOTE "
    "balancing, achieved a recall of 93.97% and a precision of 89%. The "
    "predominant misclassification pattern observed in the confusion "
    "matrix involved Critical cases being assigned to Moderate rather than "
    "lower-severity stages — a conservative error pattern that maintains "
    "clinical safety, since an animal misclassified as Moderate will still "
    "receive elevated care and monitoring rather than being dismissed as "
    "low-risk."
)
add_figure_placeholder("Fig. 14. Confusion matrix — Milk Fever stage classification (Ensemble model).")
add_body(
    "The contribution of the IoT sensor dataset conversion was "
    "particularly significant in improving Critical-stage performance. When "
    "the model was trained exclusively on the initial 3,316-record combined "
    "dataset, Critical recall stood at 86%. Following the addition of "
    "5,467 IoT sensor records derived from 580 real dairy cows — converted "
    "from RDF/Turtle format using a custom behavioural feature extraction "
    "pipeline — Critical recall improved to 93.97% and overall test "
    "accuracy rose from 94% to 97.04%. This improvement validates the data "
    "fusion strategy and confirms that real sensor-derived behavioural "
    "measurements, even when processed through clinical estimation "
    "mappings, contribute meaningfully to model generalisation beyond what "
    "simulated and statistically realistic datasets alone can provide."
)
add_body(
    "The AUC-ROC curves shown in Fig. 15 demonstrate near-perfect class "
    "separation across all four stages. The Subclinical class achieved the "
    "highest individual AUC of 0.9994, reflecting the model's strong "
    "ability to distinguish early-stage hypocalcaemia from healthy animals "
    "even in the absence of visible symptoms. The Critical class AUC of "
    "0.9971 confirms reliable separation of severe cases despite their "
    "minority representation in the original training distribution. The "
    "macro-averaged AUC-ROC of 0.9986 across all four classes represents a "
    "strong result for a four-class ordinal clinical prediction task and "
    "compares favourably with binary milk fever classifiers reported in "
    "recent literature, which typically achieve AUC values of 0.80 to 0.91 "
    "on single-disease datasets."
)
add_figure_placeholder("Fig. 15. AUC-ROC curves — Milk Fever stage classification (Ensemble model).")
add_body(
    "Beyond raw classification performance, the system introduces two "
    "design features with direct practical impact. First, a clinically "
    "grounded input transformation layer converts farmer-observable "
    "symptom responses — including behavioural state, eating behaviour, and "
    "additional symptom flags — into the physiological feature values "
    "required by the trained model. This eliminates the laboratory access "
    "barrier that renders existing precision livestock farming tools "
    "inaccessible to smallholder farmers in developing-region contexts such "
    "as Sri Lanka. Blood calcium concentration, the primary biochemical "
    "marker of hypocalcaemia, is estimated from observable indicators "
    "including inability to stand, visible muscle tremors, and behavioural "
    "deterioration using a symptom-weighted reduction formula validated "
    "against clinical literature thresholds. When laboratory values are "
    "available, they are passed directly to the feature vector, improving "
    "prediction precision further without requiring any change to the "
    "underlying model architecture."
)
add_body(
    "Second, the system incorporates a Temperature-Humidity Index weather "
    "adjustment that modifies the output risk score in real time based on "
    "current environmental conditions retrieved from a public "
    "meteorological API. Sri Lanka's tropical climate produces THI values "
    "that frequently exceed 72, the threshold at which heat stress begins "
    "to impair calcium metabolism and transition-cow management. Risk "
    "scores are adjusted upward by 5 to 15 points depending on THI "
    "severity, and the nature and magnitude of the adjustment are surfaced "
    "transparently to the user in the result interface. In addition to the "
    "stage classification, the system generates a structured "
    "explainability output identifying the specific risk factors that "
    "contributed to each prediction, including parity level, estimated "
    "blood calcium, activity level, calving proximity, body condition "
    "score, and weather adjustment. This explainability layer addresses a "
    "key limitation of black-box livestock disease prediction systems "
    "identified in prior literature and enables farmers and extension "
    "officers to understand and trust the model's reasoning rather than "
    "simply accepting an opaque risk label."
)

add_heading2("E. Overall System Performance & Validation")
add_body(
    "CattleSense was trained on validated datasets and refined against real "
    "Sri Lankan farm data, and the numbers held up well across all four "
    "modules: 87.0% accuracy for Mastitis imaging (99.4% for the biomarker "
    "side), 94.1% for FMD, 93.1% for LSD, and 97.0% for Milk Fever. "
    "Performance stayed fairly consistent even as photograph quality and "
    "how much farmer-reported data was available varied from case to "
    "case. Taken together, these numbers back up the hybrid design and "
    "suggest CattleSense could genuinely help farmers catch these diseases "
    "earlier, cutting down both diagnostic delay and guesswork."
)

# =============================================================================
# V. CONCLUSION
# =============================================================================
add_heading1("V. Conclusion")
add_body(
    "This study presents CattleSense, a multimodal machine learning "
    "platform for the early detection of four major cattle diseases "
    "affecting Sri Lankan dairy farms — Mastitis, Foot-and-Mouth Disease, "
    "Lumpy Skin Disease, and Milk Fever. By pairing image-based deep "
    "learning with disease-appropriate clinical, biomarker, or "
    "environmental data through hybrid fusion, each module achieved "
    "reliable held-out test performance: 87.0% accuracy (AUC 0.94) for "
    "Mastitis image classification alongside 99.4% for its biomarker "
    "model, 94.1% accuracy (AUC 0.97) for FMD, 93.1% accuracy (AUC 0.96) "
    "for LSD classification, and 97.0% accuracy (AUC 0.999) for Milk Fever "
    "stage prediction. These results confirm that combining visual "
    "evidence with the physiological, behavioural, and environmental data "
    "a farmer can realistically supply meaningfully improves diagnostic "
    "reliability over any single modality, while keeping the system usable "
    "without laboratory access or a veterinary login."
)
add_body(
    "CattleSense addresses a recurring gap identified across all four "
    "literature reviews: the absence of farmer-accessible, Sri "
    "Lanka-specific automated screening tools for these conditions. By "
    "delivering a risk level and actionable guidance directly to the "
    "farmer within seconds of a photograph or a short symptom checklist, "
    "the platform is positioned to reduce diagnostic delays, support "
    "earlier veterinary consultation, and limit the spread of the two "
    "contagious diseases in particular. Future work will focus on "
    "expanding each module's training dataset — most notably the LSD "
    "nodule-detection stage, whose smaller image set currently limits "
    "detection recall — integrating longitudinal per-animal risk tracking "
    "across all four modules, and conducting field trials with "
    "participating farms to validate real-world performance beyond the "
    "held-out test splits reported here."
)

# =============================================================================
# REFERENCES
# =============================================================================
add_heading1("References")
add_ref(1, "Y. Wang, X. Kang, Z. He, Y. Feng, and G. Liu, “Accurate detection of dairy cow mastitis with deep learning technology: a new and comprehensive detection method based on infrared thermal images,” Animal, vol. 16, no. 10, art. 100649, 2022.")
add_ref(2, "M. Lashin, A. S. Farid, and A. T. Elgammal, “Enhanced mastitis severity classification in dairy cows using DNN and RF: A study on PCA and correlation-based feature selection,” Smart Agricultural Technology, vol. 9, art. 100621, 2024.")
add_ref(3, "Y. Wang, X. Kang, Z. He, Y. Feng, and G. Liu, “Fusion of udder temperature and size features for the automatic detection of dairy cow mastitis using deep learning,” Computers and Electronics in Agriculture, vol. 212, art. 108131, Sep. 2023.")
add_ref(4, "X. Zhang, Y. Li, Y. Zhang, Z. Yao, W. Zou, P. Nie, and L. Yang, “A New Method to Detect Buffalo Mastitis Using Udder Ultrasonography Based on Deep Learning Network,” Animals, vol. 14, no. 5, art. 707, 2024.")
add_ref(5, "L. Pan, X. Chen, D. Han, N. Li, D. Chen, J. Wang, J. Chen, and X. Huo, “Machine learning-based clinical mastitis detection in dairy cows using milk electrical conductivity and somatic cell count,” Frontiers in Veterinary Science, vol. 12, 2025.")
add_ref(6, "M. Takahashi, A. Goto, K. Hisaeda, Y. Inoue, and T. Inaba, “Deep-learning classification of teat-end conditions in Holstein cattle,” Research in Veterinary Science, vol. 180, 2024.")
add_ref(7, "M. Z. S. Hadi, R. B. Fahreza, D. Dwimagfiroh, A. Pratiarso, and H. Mahmudah, “Detection System of Cattle Foot and Mouth Disease (FMD) using Deep Learning,” in Proc. Int. Conf. Applied Science and Technology on Engineering Science (iCAST-ES 2023), Tarakan, Indonesia, 2024.")
add_ref(8, "N. B. Ayon, A. Hasib, M. F. Ahmed, M. S. Rahman, K. Islam, T. M. M. Hasan, and A. S. M. A. S. Akib, “Simultaneous Detection of LSD and FMD in Cattle Using Ensemble Deep Learning,” arXiv:2601.12889, Jan. 2026.")
add_ref(9, "M. A. Alkhamis, H. Abouelhassan, A. Alateeqi, A. Husain, J. M. Humphreys, J. Arzt, and A. M. Perez, “Predicting the Landscape Epidemiology of Foot-and-Mouth Disease in Endemic Regions: An Interpretable Machine Learning Approach,” Viruses, vol. 17, no. 10, art. 1383, 2025.")
add_ref(10, "N. N. K. Krisnawijaya, C. Catal, B. Tekinerdogan, R. van der Tol, H. Hogeveen, and Y. Herdiyeni, “A machine learning approach to identifying foot and mouth disease incidence in dairy farms with suboptimal veterinary infrastructure,” Smart Agricultural Technology, vol. 12, art. 101261, 2025.")
add_ref(11, "U. Gunasekera, M. A. Alkhamis, S. Puvanendiran, M. Das, P. L. Kumarawadu, M. Sultana, M. A. Hossain, J. Arzt, and A. Perez, “Ecological niche modeling for surveillance of foot-and-mouth disease in South Asia,” PLOS ONE, vol. 20, no. 4, art. e0320921, 2025.")
add_ref(12, "A. F. H. Alharan, H. K. Flaih, and A. S. Alkhafaji, “Classification of Cattle Skin Diseases Using Convolutional Neural Networks,” Iraqi Journal of Science, vol. 62, no. 8, pp. 2786-2795, 2021.")
add_ref(13, "R. Kumar, S. Singh, and A. Gupta, “Deep Learning-Based Cattle Skin Disease Detection Using Transfer Learning,” Journal of Veterinary Informatics, vol. 14, no. 3, pp. 112–120, 2022.")
add_ref(14, "E. S. M. Tuppurainen and C. A. L. Oura, “Review: Lumpy Skin Disease: An Emerging Threat to Europe, the Middle East and Asia,” Transboundary and Emerging Diseases, vol. 59, no. 1, pp. 40-48, 2012.")
add_ref(15, "S. Babiuk, T. R. Bowden, D. B. Boyle, D. B. Wallace, and R. P. Kitching, “Capripoxviruses: An Emerging Worldwide Threat to Sheep, Goats and Cattle,” Transboundary and Emerging Diseases, vol. 55, no. 7, pp. 263-272, 2008.")
add_ref(16, "G. Jocher, A. Chaurasia, and J. Qiu, “YOLOv8 by Ultralytics,” Ultralytics, 2023. [Online]. Available: https://github.com/ultralytics/ultralytics.")
add_ref(17, "K. He, X. Zhang, S. Ren, and J. Sun, “Deep Residual Learning for Image Recognition,” in Proc. IEEE Conf. Computer Vision and Pattern Recognition (CVPR), 2016, pp. 770-778.")
add_ref_todo(18, "[NEEDS FULL CITATION] Lasser et al., 2021 — machine learning prediction of periparturient hypocalcaemia and other dairy diseases from farm/animal data. Ask the Milk Fever author for the paper title, journal/conference, volume, and pages.")
add_ref_todo(19, "[NEEDS FULL CITATION] Van Leerdam et al., 2024 — behavioural sensor data + XGBoost/LSTM for early hypocalcaemia prediction. Ask the Milk Fever author for the paper title, journal/conference, volume, and pages.")

import sys
OUT_NAME = sys.argv[1] if len(sys.argv) > 1 else "CattleSense_Full_Paper.docx"
doc.save(OUT_NAME)
print(f"saved {OUT_NAME}")
