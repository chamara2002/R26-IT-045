"""Builds the condensed, one-page Word version of the LSD Results and Discussion.

Run with the lumpy-module venv (needs python-docx):
    venv/Scripts/python.exe docs/generate_results_doc_onepage.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIG_DIR = "figures"

doc = Document()

# ---- Page setup: narrow margins, two columns (IEEE-style) -----------------
section = doc.sections[0]
section.top_margin = Mm(15)
section.bottom_margin = Mm(15)
section.left_margin = Mm(14)
section.right_margin = Mm(14)

sectPr = section._sectPr
cols = OxmlElement("w:cols")
cols.set(qn("w:num"), "2")
cols.set(qn("w:space"), "425")  # ~0.3in gutter
sectPr.append(cols)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(9.5)
style.paragraph_format.space_after = Pt(5)
style.paragraph_format.line_spacing = 1.05


def add_para(text, space_after=5, size=9.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    parts = text.split("**")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.size = Pt(size)
        if i % 2 == 1:
            run.bold = True
    return p


def set_cell_shading(cell, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(caption, headers, rows):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(3)
    run = cap.add_run(caption)
    run.bold = True
    run.font.size = Pt(8.5)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "1F4E79")
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(255, 255, 255)

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_figure(path, caption, width_in=3.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(f"{FIG_DIR}/{path}", width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(8)
    cap.paragraph_format.space_after = Pt(6)


# ---- Title (spans both columns via a separate one-column section block) ---
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("IV. RESULTS AND DISCUSSION")
r.font.size = Pt(12.5)
r.font.bold = True
title_p.paragraph_format.space_after = Pt(8)

# ---- Body ------------------------------------------------------------------
add_para(
    "**Detection.** The YOLOv8s lesion detector was trained for up to 150 epochs "
    "(640×640, lr0=0.001) on 546 Roboflow-annotated images (129 background "
    "negatives) and evaluated on a held-out set of 24 images (261 nodule instances). "
    "Training plateaued and began to overfit after roughly epoch 50, with validation "
    "box loss turning upward even as the training curve kept falling. The best "
    "checkpoint reached mAP@0.5 = 0.1383 on validation; on the held-out test set this "
    "fell to 0.1183 — well short of the project’s own 80% target — with "
    "precision 0.19 and recall 0.28 (Table I)."
)

add_para(
    "**Classification.** ResNet50 was fine-tuned in two phases (head-only, then "
    "partial backbone unfreezing) on 382 whole photographs (320 positive, 62 healthy), "
    "split 266/58/58. On the 58-image test set it reached 93.1% accuracy and an AUC of "
    "0.9637, correctly classifying 48 of 49 positive and 6 of 9 healthy cases. Training "
    "accuracy and AUC saturated at 1.0 well before validation performance stabilised "
    "near 95%, and a held-out example returned a raw probability of 100.0% — a mild "
    "but real overconfidence traceable to the 320:62 class imbalance, which is why the "
    "deployed app never exposes this raw number directly."
)

add_para(
    "**Integrated pipeline.** Chaining detection and classification (region score = "
    "detection confidence × classification probability, image score = max across "
    "regions, 30% decision threshold) on the same aligned 24-image test set dropped "
    "overall accuracy to 41.67%, with 100% precision but a 77.78% false-negative rate "
    "(Fig. 1, Table I). This is not a joint failure: because the image score is gated "
    "on YOLOv8 finding at least one region first, and the detector’s own recall on "
    "this set was only 27.59%, most true-positive images never reach the (individually "
    "strong) classifier at all. Average processing time, 2.69 s/image, comfortably met "
    "the 10–15 s target."
)

add_table(
    "TABLE I. HEADLINE PERFORMANCE ACROSS ALL THREE STAGES",
    ["Stage", "Key metrics", "Test set"],
    [
        ["YOLOv8s Detection", "mAP@0.5=11.83%, P=0.19, R=0.28", "n=24 (261 inst.)"],
        ["ResNet50 Classification", "Accuracy=93.10%, AUC=0.9637", "n=58"],
        ["Integrated Pipeline", "Accuracy=41.67%, FPR=0.00%, FNR=77.78%", "n=24 (aligned)"],
    ],
)

add_figure(
    "integrated_cell14_out0.png",
    "Fig. 1. End-to-end pipeline confusion matrix, aligned held-out test.",
    width_in=2.9,
)

add_para(
    "**Discussion.** The bottleneck is squarely the detector, not the classifier, "
    "which points future effort toward expanding and re-auditing the nodule-detection "
    "dataset rather than further classifier tuning. The current precision/recall "
    "balance — essentially never a false alarm, but missing most real cases — "
    "is a defensible starting posture for a farmer-facing screening tool, and is the "
    "direct motivation for the clinical symptom-fusion layer used in deployment "
    "(Section III), intended to catch exactly the visually-quiet cases this evaluation "
    "shows the detector currently misses."
)

doc.save("LSD_Results_and_Discussion_OnePage.docx")
print("saved LSD_Results_and_Discussion_OnePage.docx")
